"""DOCX processor — .docx files (requires python-docx)."""

import io
from typing import Any

from app.ai.document_processors.base import (
    DocumentProcessor,
    ProcessedChunk,
    ProcessedDocument,
    ProcessorError,
)


class DOCXProcessor(DocumentProcessor):
    """Process Microsoft Word .docx files using python-docx.

    Extracts text paragraph-by-paragraph. Heading styles (Heading 1, 2, 3)
    are detected and used to build the `heading_path` for each chunk.
    """

    @property
    def supported_formats(self) -> list[str]:
        return ["docx"]

    async def process_bytes(
        self,
        data: bytes,
        filename: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> ProcessedDocument:
        meta = dict(metadata or {})
        try:
            from docx import Document
        except ImportError as e:  # pragma: no cover — optional dep
            raise ProcessorError(
                "python-docx not installed. Run: pip install python-docx"
            ) from e

        try:
            doc = Document(io.BytesIO(data))
        except Exception as e:
            raise ProcessorError(f"Cannot read DOCX: {e}") from e

        # Extract title from core properties
        title = meta.get("title")
        if not title:
            try:
                cp = doc.core_properties
                if cp.title:
                    title = cp.title
            except Exception:
                pass
        if not title:
            title = filename or "Untitled DOCX"

        chunks: list[ProcessedChunk] = []
        full_text_parts: list[str] = []
        current_heading_path: list[str] = []
        current_section_lines: list[str] = []
        position = 0

        def flush() -> None:
            nonlocal position
            if not current_section_lines:
                return
            section_text = "\n".join(current_section_lines).strip()
            if not section_text:
                current_section_lines.clear()
                return
            chunks.append(
                ProcessedChunk(
                    text=section_text,
                    heading_path=list(current_heading_path),
                    position=position,
                )
            )
            position += 1
            current_section_lines.clear()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            full_text_parts.append(text)
            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name.startswith("heading"):
                # Flush previous section
                flush()
                # Parse heading level
                try:
                    level = int(style_name.replace("heading", "").strip())
                except ValueError:
                    level = 1
                current_heading_path = current_heading_path[: level - 1]
                current_heading_path.append(text)
                current_section_lines.append(text)
            else:
                current_section_lines.append(text)
                # If section is getting too large, flush
                joined = "\n".join(current_section_lines)
                if len(joined) > 1500:
                    flush()

        flush()

        full_text = "\n\n".join(full_text_parts)
        return ProcessedDocument(
            text=full_text,
            chunks=chunks,
            title=title,
            language=meta.get("language", "en"),
            metadata={
                **meta,
                "paragraph_count": len(doc.paragraphs),
            },
        )

    async def process_text(
        self,
        text: str,
        filename: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> ProcessedDocument:
        raise ProcessorError(
            "DOCXProcessor cannot process plain text — use process_bytes() with raw DOCX data"
        )
